import os
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_table_header_bg_color(cell, color_hex):
    """
    Set background shading for a given cell.
    """
    tblCell = cell._tc
    tblCellProperties = tblCell.get_or_add_tcPr()
    clShading = OxmlElement('w:shd')
    clShading.set(qn('w:fill'), color_hex)  # Hex of desired color
    tblCellProperties.append(clShading)
    return cell



def create_permit_document(regulator, airline, representative_name, start_date, end_date, dataframe):
    document = Document()

    # Decrease white border margins on all sides
    section = document.sections[0]
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)

    header_paragraph = document.sections[0].header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    header_run = header_paragraph.add_run()
    header_run.add_picture(os.path.join('visual_elements', 'header.png'), height=Cm(1.9))

    footer_paragraph = document.sections[0].footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer_paragraph.add_run()
    footer_run.add_picture(os.path.join('visual_elements', 'footer.png'), width=Cm(18.5))

    document.styles['Normal'].font.name = 'Segoe UI'
    document.styles['Normal'].font.size = Pt(11)

    document.add_paragraph(f"Dear {regulator},")
    document.add_paragraph(f"On behalf of {airline}, I am writing to request your kind consideration for landing approval to operate scheduled commercial flights between the dates of {start_date} to {end_date}.")
    document.add_paragraph("We have submitted our proposed flight schedule below for your review and approval:")

    # Remove the last two columns from the dataframe
    dataframe = dataframe.iloc[:, :-2]

    table = document.add_table(rows=1, cols=len(dataframe.columns))
    header_cells = table.rows[0].cells
    header_fill_color_hex = "001A23"  # Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}

    for idx, column_name in enumerate(dataframe.columns):
        header_cells[idx].text = column_name
        set_table_header_bg_color(header_cells[idx], header_fill_color_hex)


    # Add data rows
    for index, row in dataframe.iterrows():
        cells = table.add_row().cells
        for idx, col_name in enumerate(dataframe.columns):
            cells[idx].text = str(row[col_name])
            cells[idx].paragraphs[0].runs[0].font.size = Pt(9.5)

    # Leave more blank space under the table
    # document.add_paragraph("\n")
    document.add_paragraph("We assure you that our operations will comply with all the regulatory requirements, and we will work closely with your team to ensure that all necessary procedures are followed.")
    document.add_paragraph("We appreciate your time and attention to our request, and we look forward to hearing back from you soon. Please do not hesitate to contact us if you require any further information.")
    document.add_paragraph(f"Thank you for your kind consideration.")
    document.add_paragraph(f"Sincerely,")

    # Leave less blank space between the representative_name and airline
    rep_name_paragraph = document.add_paragraph()
    rep_name_run = rep_name_paragraph.add_run(f"{representative_name}")
    rep_name_run.add_break()
    rep_name_run.add_text(f"{airline}")

    return document

def save_permit_document(document, filename):
    document.save(filename)

